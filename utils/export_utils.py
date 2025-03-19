import os
import io
from typing import Optional, Dict, Any
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_pdf(transcription: Dict[str, Any], output_path: Optional[str] = None) -> str:
    """
    Generate a PDF file from a transcription.
    
    Args:
        transcription: Transcription dictionary from the database
        output_path: Path to save the PDF file (optional)
        
    Returns:
        Path to the generated PDF file
    """
    # If no output path is provided, create one in the current directory
    if output_path is None:
        filename = transcription['original_filename'].rsplit('.', 1)[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"{filename}_{timestamp}.pdf"
    
    # Create a PDF buffer
    buffer = io.BytesIO()
    
    # Create the PDF document
    doc = SimpleDocTemplate(buffer, pagesize=letter, title=f"Transcription - {transcription['original_filename']}")
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = styles['Title']
    heading_style = styles['Heading2']
    normal_style = styles['Normal']
    
    # Custom style for transcription text
    transcription_style = ParagraphStyle(
        'TranscriptionStyle',
        parent=normal_style,
        spaceBefore=6,
        spaceAfter=6,
        leading=14
    )
    
    # Create document content
    content = []
    
    # Add title
    content.append(Paragraph(f"Transcription: {transcription['original_filename']}", title_style))
    content.append(Spacer(1, 12))
    
    # Add metadata table
    metadata = [
        ['File Name:', transcription['original_filename']],
        ['Duration:', f"{transcription['duration_seconds']:.2f} seconds"],
        ['File Type:', transcription['file_type']],
        ['Date:', transcription['created_at']]
    ]
    
    metadata_table = Table(metadata, colWidths=[100, 350])
    metadata_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    
    content.append(metadata_table)
    content.append(Spacer(1, 20))
    
    # Add processed transcription
    content.append(Paragraph("Processed Transcription:", heading_style))
    content.append(Spacer(1, 6))
    
    # Split the processed transcription into paragraphs and add each as a Paragraph
    for paragraph in transcription['processed_transcription'].split('\n\n'):
        if paragraph.strip():
            content.append(Paragraph(paragraph, transcription_style))
            content.append(Spacer(1, 6))
    
    # Build the PDF document
    doc.build(content)
    
    # Get the PDF content from the buffer
    pdf_content = buffer.getvalue()
    buffer.close()
    
    # Write the PDF to a file
    with open(output_path, 'wb') as f:
        f.write(pdf_content)
    
    return output_path

def export_plaintext(transcription: Dict[str, Any], output_path: Optional[str] = None) -> str:
    """
    Export a transcription as plaintext.
    
    Args:
        transcription: Transcription dictionary from the database
        output_path: Path to save the text file (optional)
        
    Returns:
        Path to the generated text file
    """
    # If no output path is provided, create one in the current directory
    if output_path is None:
        filename = transcription['original_filename'].rsplit('.', 1)[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"{filename}_{timestamp}.txt"
    
    # Create the text content
    content = f"Transcription: {transcription['original_filename']}\n"
    content += f"Duration: {transcription['duration_seconds']:.2f} seconds\n"
    content += f"File Type: {transcription['file_type']}\n"
    content += f"Date: {transcription['created_at']}\n\n"
    content += f"Processed Transcription:\n\n"
    content += transcription['processed_transcription']
    
    # Write the text to a file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    return output_path

def export_to_word(transcription: Dict[str, Any], output_path: Optional[str] = None) -> str:
    """
    Export a transcription as a Microsoft Word document.
    
    Args:
        transcription: Transcription dictionary from the database
        output_path: Path to save the Word file (optional)
        
    Returns:
        Path to the generated Word file
    """
    # If no output path is provided, create one in the current directory
    if output_path is None:
        filename = transcription['original_filename'].rsplit('.', 1)[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"{filename}_{timestamp}.docx"
    
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading(f"Transcription: {transcription['original_filename']}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata section
    doc.add_heading('File Information', level=2)
    
    # Create 2x2 table for metadata
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    # Add metadata to the table
    cells = table.rows[0].cells
    cells[0].text = 'File Name:'
    cells[1].text = transcription['original_filename']
    
    cells = table.rows[1].cells
    cells[0].text = 'Duration:'
    cells[1].text = f"{transcription['duration_seconds']:.2f} seconds"
    
    cells = table.rows[2].cells
    cells[0].text = 'File Type:'
    cells[1].text = transcription['file_type']
    
    cells = table.rows[3].cells
    cells[0].text = 'Date:'
    cells[1].text = str(transcription['created_at'])
    
    # Format table
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    
    # Add some space
    doc.add_paragraph()
    
    # Add transcription section
    doc.add_heading('Processed Transcription', level=2)
    
    # Split the transcription into paragraphs and add each as a separate paragraph
    for paragraph_text in transcription['processed_transcription'].split('\n\n'):
        if paragraph_text.strip():
            p = doc.add_paragraph(paragraph_text)
            # Set line spacing
            p.paragraph_format.line_spacing = 1.15
            # Add some space after each paragraph
            p.paragraph_format.space_after = Pt(8)
    
    # Save the document
    doc.save(output_path)
    
    return output_path
